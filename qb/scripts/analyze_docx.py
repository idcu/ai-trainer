import docx
import re
import os

def analyze_docx(file_path, question_type):
    doc = docx.Document(file_path)
    all_text = []
    
    print(f"\n{'='*60}")
    print(f"分析文件: {os.path.basename(file_path)}")
    print(f"{'='*60}")
    
    # 收集所有段落文本
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append((i, text))
    
    print(f"总段落数（含空行）: {len(doc.paragraphs)}")
    print(f"有效段落数: {len(all_text)}")
    
    # 合并所有文本
    full_text = '\n'.join([text for _, text in all_text])
    
    # 查找题目编号
    num_pattern = re.compile(r'(\d+)[、\.]')
    nums = []
    for match in num_pattern.finditer(full_text):
        nums.append(int(match.group(1)))
    
    if nums:
        print(f"\n检测到的题目编号范围: {min(nums)} - {max(nums)}")
        print(f"期望的题目数量: {max(nums)}")
        
        # 检查缺失的编号
        missing = []
        for i in range(min(nums), max(nums)+1):
            if i not in nums:
                missing.append(i)
        if missing:
            print(f"缺失的编号: {missing}")
        
        # 检查重复的编号
        from collections import Counter
        num_counts = Counter(nums)
        duplicates = {k: v for k, v in num_counts.items() if v > 1}
        if duplicates:
            print(f"重复的编号: {duplicates}")
    
    # 查找正确答案
    ans_pattern = re.compile(r'正确答案[：:]\s*([正确错误A-D]+)')
    answers = []
    for match in ans_pattern.finditer(full_text):
        answers.append(match.group(1))
    
    print(f"\n检测到的答案数量: {len(answers)}")
    
    # 查找解析
    analysis_pattern = re.compile(r'解析[：:]')
    analyses = []
    for match in analysis_pattern.finditer(full_text):
        analyses.append(match.start())
    
    print(f"检测到的解析数量: {len(analyses)}")
    
    # 显示前几个段落的内容
    print(f"\n前20个有效段落预览:")
    print("-" * 60)
    for i, (para_idx, text) in enumerate(all_text[:20]):
        print(f"[{i+1}] 段落{para_idx}: {text[:80]}..." if len(text) > 80 else f"[{i+1}] 段落{para_idx}: {text}")
    
    return all_text

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    qb_dir = os.path.dirname(base_dir)
    project_dir = os.path.dirname(qb_dir)
    word_dir = os.path.join(project_dir, "word")
    
    files = [
        ("人工智能训练师理论复习（判断汇总）.docx", "judge"),
        ("人工智能训练师理论复习（单选汇总）.docx", "single"),
        ("人工智能训练师理论复习（多选汇总）.docx", "multi")
    ]
    
    for filename, qtype in files:
        filepath = os.path.join(word_dir, filename)
        if os.path.exists(filepath):
            analyze_docx(filepath, qtype)

if __name__ == "__main__":
    main()
